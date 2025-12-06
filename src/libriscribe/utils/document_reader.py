# src/libriscribe/utils/document_reader.py

import logging
from pathlib import Path
from typing import Optional
import re

logger = logging.getLogger(__name__)


class DocumentReader:
    """
    Universal document reader supporting multiple formats.
    
    Supports: .txt, .md, .rtf, .docx, .pdf
    """
    
    def __init__(self):
        self.supported_formats = ['.txt', '.md', '.rtf', '.docx', '.pdf']
    
    def read_document(self, file_path: Path) -> Optional[str]:
        """
        Read a document and return its text content.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Text content or None if failed
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return None
        
        suffix = file_path.suffix.lower()
        
        try:
            if suffix in ['.txt', '.md']:
                return self._read_text(file_path)
            elif suffix == '.rtf':
                return self._read_rtf(file_path)
            elif suffix == '.docx':
                return self._read_docx(file_path)
            elif suffix == '.pdf':
                return self._read_pdf(file_path)
            else:
                logger.warning(f"Unsupported format: {suffix}")
                return None
                
        except Exception as e:
            logger.exception(f"Error reading {file_path}: {e}")
            return None
    
    def _read_text(self, file_path: Path) -> str:
        """Read plain text or markdown file."""
        return file_path.read_text(encoding='utf-8')
    
    def _read_rtf(self, file_path: Path) -> str:
        """Read RTF file and extract text."""
        try:
            # Try using striprtf library if available
            from striprtf.striprtf import rtf_to_text
            content = file_path.read_text(encoding='utf-8')
            return rtf_to_text(content)
        except ImportError:
            # Fallback: basic RTF stripping
            logger.warning("striprtf not installed, using basic RTF parser")
            return self._basic_rtf_strip(file_path)
    
    def _basic_rtf_strip(self, file_path: Path) -> str:
        """Basic RTF text extraction (fallback)."""
        content = file_path.read_text(encoding='utf-8', errors='ignore')
        
        # Remove RTF control words and groups
        # This is a simple approach - may not be perfect
        text = re.sub(r'\\[a-z]+\d*\s?', '', content)  # Remove control words
        text = re.sub(r'[{}]', '', text)  # Remove braces
        text = re.sub(r'\\\'[0-9a-f]{2}', '', text)  # Remove hex codes
        text = re.sub(r'\\\*.*?;', '', text)  # Remove special groups
        
        # Clean up whitespace
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        return text
    
    def _read_docx(self, file_path: Path) -> str:
        """Read DOCX file and extract text and images."""
        try:
            from docx import Document
            import zipfile
            import base64
            from io import BytesIO
            
            doc = Document(file_path)
            
            # Extract text from all paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract and analyze images
            image_descriptions = self._extract_docx_images(file_path)
            
            # Combine text and image descriptions
            full_text = '\n\n'.join(text_parts)
            
            if image_descriptions:
                full_text += "\n\n--- EMBEDDED IMAGES ---\n\n"
                full_text += '\n\n'.join(image_descriptions)
            
            return full_text
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return None
    
    def _extract_docx_images(self, file_path: Path) -> list:
        """Extract images from DOCX and get descriptions using vision LLM."""
        try:
            from docx import Document
            import zipfile
            from PIL import Image
            from io import BytesIO
            
            # Try to get LLM client for vision analysis
            try:
                from libriscribe.utils.llm_client import LLMClient
                from libriscribe.settings import settings
                llm_client = LLMClient(settings)
                has_vision = True
            except:
                has_vision = False
                logger.warning("LLM client not available for image analysis")
            
            doc = Document(file_path)
            descriptions = []
            
            # Extract images from document relationships
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_data = rel.target_part.blob
                        
                        if has_vision:
                            # Analyze image with vision LLM
                            description = self._analyze_image_with_llm(image_data, llm_client)
                            if description:
                                descriptions.append(f"Image: {description}")
                        else:
                            descriptions.append(f"Image: [Image found but not analyzed - vision LLM not available]")
                    
                    except Exception as e:
                        logger.warning(f"Could not process image: {e}")
            
            return descriptions
            
        except Exception as e:
            logger.warning(f"Could not extract images: {e}")
            return []
    
    def _analyze_image_with_llm(self, image_data: bytes, llm_client) -> str:
        """Analyze image using vision LLM."""
        try:
            import base64
            from io import BytesIO
            from PIL import Image
            
            # Convert to base64 for LLM
            image_base64 = base64.b64encode(image_data).decode('utf-8')
            
            # Use vision-capable model (Gemini supports images)
            prompt = """Describe this chart/image from an editorial report. 
            Focus on:
            - What type of chart/graph it is
            - Key data points or trends shown
            - What it indicates about the manuscript
            
            Be concise and specific."""
            
            # Note: This requires vision-capable LLM
            # Gemini 2.0 Flash supports images
            response = llm_client.generate_content_with_image(
                prompt,
                image_base64,
                max_tokens=300
            )
            
            return response.strip()
            
        except Exception as e:
            logger.warning(f"Could not analyze image with LLM: {e}")
            return "[Image analysis failed]"

    
    def _read_pdf(self, file_path: Path) -> str:
        """Read PDF file and extract text and images."""
        try:
            import PyPDF2
            
            text_parts = []
            
            # Extract text from pages
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"--- Page {page_num} ---\n{page_text}")
            
            full_text = '\n\n'.join(text_parts)
            
            # Extract and analyze images from PDF
            image_descriptions = self._extract_pdf_images(file_path)
            
            if image_descriptions:
                full_text += "\n\n--- EMBEDDED IMAGES/CHARTS ---\n\n"
                full_text += '\n\n'.join(image_descriptions)
            
            return full_text
            
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            return None
    
    def _extract_pdf_images(self, file_path: Path) -> list:
        """Extract images from PDF pages and analyze with vision LLM."""
        try:
            from pdf2image import convert_from_path
            import base64
            from io import BytesIO
            
            # Try to get LLM client for vision analysis
            try:
                from libriscribe.utils.llm_client import LLMClient
                from libriscribe.settings import settings
                llm_client = LLMClient(settings)
                has_vision = True
            except:
                has_vision = False
                logger.warning("LLM client not available for image analysis")
            
            descriptions = []
            
            # Convert PDF pages to images
            # Note: This requires poppler to be installed on Windows
            try:
                images = convert_from_path(file_path, dpi=150)
                
                for page_num, image in enumerate(images, 1):
                    if has_vision:
                        # Convert PIL image to base64
                        buffered = BytesIO()
                        image.save(buffered, format="PNG")
                        image_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                        
                        # Analyze page image for charts/graphs
                        description = self._analyze_pdf_page_for_charts(image_base64, page_num, llm_client)
                        if description and description != "[No charts found]":
                            descriptions.append(f"Page {page_num}: {description}")
                
            except Exception as e:
                logger.warning(f"Could not convert PDF to images (poppler may not be installed): {e}")
                logger.info("Install poppler for Windows: https://github.com/oschwartz10612/poppler-windows/releases/")
            
            return descriptions
            
        except ImportError as e:
            logger.warning(f"pdf2image not installed: {e}")
            return []
    
    def _analyze_pdf_page_for_charts(self, image_base64: str, page_num: int, llm_client) -> str:
        """Analyze PDF page image for charts and graphs."""
        try:
            prompt = """Analyze this PDF page for charts, graphs, or visual data.

If you find any charts/graphs:
- Describe what type (bar chart, line graph, pie chart, etc.)
- Extract key data points and trends
- Explain what it shows about the manuscript

If there are NO charts/graphs on this page, respond with: [No charts found]

Be concise and focus only on visual data, not regular text."""
            
            response = llm_client.generate_content_with_image(
                prompt,
                image_base64,
                max_tokens=400
            )
            
            return response.strip()
            
        except Exception as e:
            logger.warning(f"Could not analyze PDF page {page_num}: {e}")
            return "[Image analysis failed]"

    
    def read_folder(self, folder_path: Path) -> dict:
        """
        Read all supported documents in a folder.
        
        Args:
            folder_path: Path to folder
            
        Returns:
            Dict of {filename: content}
        """
        folder_path = Path(folder_path)
        documents = {}
        
        if not folder_path.exists():
            logger.error(f"Folder not found: {folder_path}")
            return documents
        
        for file_path in folder_path.iterdir():
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                content = self.read_document(file_path)
                if content:
                    documents[file_path.name] = content
                    logger.info(f"Read {file_path.name} ({len(content)} chars)")
        
        return documents
